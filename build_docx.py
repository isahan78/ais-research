"""Render COMBINED_GOOGLEDOC.md to a polished .docx with the two figures inline.

Handles the constructs in the deliverable: #/##/### headings, **bold**,
*italic* / _italic_, `code`, [text](url) links, numbered/bullet lists (with
hard-wrapped continuation lines joined), pipe tables, whole-line _italic_
subtitles, blockquote captions, --- rules (page break before 'Part 2'), and the
[INSERT FIGURE ... `path`] markers (replaced by the image + caption).

Design: editorial, not default-Word. Georgia body, deep-navy headings with an
accent rule under the title and part headings, a custom clean results table
(navy header, alternating row tint, hairline borders), and a page-numbered
footer. Opens natively in Google Docs / Word.

Run: python build_docx.py  ->  The_Probe_Text_Gap.docx
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
SRC = HERE / "COMBINED_GOOGLEDOC.md"
OUT = HERE / "The_Probe_Text_Gap.docx"

NAVY = RGBColor(0x15, 0x2A, 0x45)      # headings / title
ACCENT = RGBColor(0x2C, 0x5F, 0x8A)    # part headings, links
GREY = RGBColor(0x5A, 0x63, 0x6E)      # captions, subtitle
NAVY_HEX, TINT_HEX, RULE_HEX = "152A45", "EEF2F7", "C9D4E0"

INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\w)_.+?_(?!\w)|(?<!\*)\*[^*]+?\*(?!\*)|`.+?`|\[.+?\]\(.+?\))")
FIG = re.compile(r"INSERT FIGURE.*?`([^`]+)`")
LINK = re.compile(r"\[(.+?)\]\((.+?)\)")
SUBTITLE = re.compile(r"^_[^_].*_$")


def add_runs(p, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif (tok.startswith("_") and tok.endswith("_")) or (tok.startswith("*") and tok.endswith("*")):
            p.add_run(tok[1:-1]).italic = True
        elif tok.startswith("[") and LINK.match(tok):
            m = LINK.match(tok)
            r = p.add_run(m.group(1)); r.font.color.rgb = ACCENT; r.underline = True
        else:
            p.add_run(tok)


def bottom_border(paragraph, color=NAVY_HEX, sz=8, space=6):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", str(sz)), ("w:space", str(space)), ("w:color", color)):
        b.set(qn(k), v)
    pbdr.append(b); pPr.append(pbdr)


def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_); tcPr.append(shd)


def table_borders(table, color=RULE_HEX, sz=4):
    tblPr = table._tbl.tblPr; borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        e = OxmlElement(f"w:{edge}")
        for k, v in (("w:val", "single"), ("w:sz", str(sz)), ("w:space", "0"), ("w:color", color)):
            e.set(qn(k), v)
        borders.append(e)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def page_footer(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Isahan Khan · The Probe–Text Gap · page "); r.font.size = Pt(8); r.font.color.rgb = GREY
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    rr = p.add_run(); rr.font.size = Pt(8); rr.font.color.rgb = GREY
    rr._r.append(fld1); rr._r.append(instr); rr._r.append(fld2)


def setup_styles(doc):
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.05)
        s.top_margin = s.bottom_margin = Inches(0.95)
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"; normal.font.size = Pt(10.5)
    pf = normal.paragraph_format; pf.line_spacing = 1.16; pf.space_after = Pt(7); pf.space_before = Pt(0)

    def h(name, size, color, before, after=4, rule=False):
        st = doc.styles[name]; f = st.font
        f.name = "Georgia"; f.size = Pt(size); f.bold = True; f.color.rgb = color
        p = st.paragraph_format; p.space_before = Pt(before); p.space_after = Pt(after); p.keep_with_next = True
        return st

    h("Title", 23, NAVY, 0, 2)
    h("Heading 1", 15.5, ACCENT, 20, 6)
    h("Heading 2", 12.5, NAVY, 14, 4)
    h("Heading 3", 11, NAVY, 10, 3)
    for ln in ("List Bullet", "List Number"):
        lp = doc.styles[ln].paragraph_format; lp.space_after = Pt(4); lp.line_spacing = 1.14


def add_caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(12)
    add_runs(p, text.strip())
    for r in p.runs:
        r.font.size = Pt(8.5); r.font.color.rgb = GREY; r.italic = True


def add_subtitle(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10); r.font.color.rgb = GREY


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], cells[2:]
    t = doc.add_table(rows=1, cols=len(header)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for j, htext in enumerate(header):
        c = t.rows[0].cells[j]; shade(c, NAVY_HEX); c.paragraphs[0].clear()
        add_runs(c.paragraphs[0], htext)
        for r in c.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(body):
        wc = t.add_row().cells
        for j, val in enumerate(row):
            if j < len(wc):
                if i % 2 == 1:
                    shade(wc[j], TINT_HEX)
                wc[j].paragraphs[0].clear(); add_runs(wc[j].paragraphs[0], val)
                for r in wc[j].paragraphs[0].runs:
                    r.font.size = Pt(8.5)
    table_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def flush_para(doc, buf, style):
    if not buf:
        return
    text = " ".join(buf).strip()
    if text:
        add_runs(doc.add_paragraph(style=style), text)
    buf.clear()


def main():
    doc = Document()
    setup_styles(doc)
    page_footer(doc)

    lines = SRC.read_text().splitlines()
    buf, list_style = [], None
    i, first_h1, just_title = 0, True, False
    while i < len(lines):
        line = lines[i].rstrip(); stripped = line.strip()

        if FIG.search(line):
            flush_para(doc, buf, list_style); list_style = None
            path = HERE / FIG.search(line).group(1)
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.1))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
            i += 1; cap = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                cap.append(lines[i].strip().lstrip(">").strip()); i += 1
            if cap:
                add_caption(doc, " ".join(cap))
            continue

        if not stripped:
            flush_para(doc, buf, list_style); list_style = None; i += 1; continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_para(doc, buf, list_style); list_style = None
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i]); i += 1
            add_table(doc, tbl); continue

        if stripped == "---":
            flush_para(doc, buf, list_style); list_style = None
            if i + 2 < len(lines) and "Part 2" in lines[i + 2]:
                doc.add_page_break()
            i += 1; continue

        m = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if m:
            flush_para(doc, buf, list_style); list_style = None
            level, txt = len(m.group(1)), m.group(2)
            if level == 1 and first_h1:
                p = doc.add_heading("", level=0); add_runs(p, txt); bottom_border(p)
                first_h1, just_title = False, True
            else:
                p = doc.add_heading("", level=level); add_runs(p, txt)
                if level == 1:
                    bottom_border(p, color=RULE_HEX, sz=6)
                just_title = False
            i += 1; continue

        if just_title and SUBTITLE.match(stripped):
            add_subtitle(doc, stripped); i += 1; continue
        just_title = False

        mnum = re.match(r"^(\d+)\.\s+(.*)", stripped)
        mbul = re.match(r"^[-*]\s+(.*)", stripped)
        if mnum:
            flush_para(doc, buf, list_style); buf, list_style = [mnum.group(2)], "List Number"; i += 1; continue
        if mbul:
            flush_para(doc, buf, list_style); buf, list_style = [mbul.group(1)], "List Bullet"; i += 1; continue

        buf.append(stripped); i += 1

    flush_para(doc, buf, list_style)
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
